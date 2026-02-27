from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import FileResponse, Response
from sqlalchemy.orm import Session
from typing import List
from backend.database import get_db, SessionLocal
from backend.models import Document
from backend.schemas import DocumentResponse
import os
import asyncio
import PyPDF2
from docx import Document as DocxDocument
from urllib.parse import quote
import logging

router = APIRouter()
logger = logging.getLogger(__name__)

# Temporary upload directory (in production, use object storage)
# Use project-based uploads directory for cross-platform compatibility
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

async def extract_content(file_path: str, file_type: str) -> str:
    """Extract content from different file types"""
    content = ""
    
    try:
        if file_type == "application/pdf":
            # Extract content from PDF
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    content += page.extract_text() + "\n"
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Extract content from DOCX
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
        
        elif file_type in ["text/plain", "text/markdown", "text/html"]:
            # Extract content from text files
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        
        logger.info(f"Successfully extracted {len(content)} characters from {file_path}")
    
    except Exception as e:
        logger.error(f"Error extracting content from {file_path}: {e}")
        content = ""
    
    return content

async def process_document(document_id: int, file_path: str, file_type: str):
    """Process document asynchronously with its own database session"""
    db = SessionLocal()
    try:
        # Update status to processing
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            logger.error(f"Document {document_id} not found in database")
            return
        
        document.status = "processing"
        db.commit()
        logger.info(f"Started processing document {document_id}: {document.filename}")
        
        # Extract content
        content = await extract_content(file_path, file_type)
        
        # Update document with content and status
        document.content = content
        document.status = "completed" if content else "failed"
        db.commit()
        
        logger.info(f"Document {document_id} processing completed with status: {document.status}, content length: {len(content) if content else 0}")
    
    except Exception as e:
        logger.error(f"Error processing document {document_id}: {e}")
        try:
            document = db.query(Document).filter(Document.id == document_id).first()
            if document:
                document.status = "failed"
                db.commit()
        except Exception as commit_error:
            logger.error(f"Error updating document status: {commit_error}")
    
    finally:
        db.close()

@router.post("/upload/{notebook_id}", response_model=DocumentResponse)
async def upload_document(
    notebook_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Upload a document to a notebook"""
    
    # Validate file type
    allowed_types = [
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
        "text/markdown",
        "text/html"
    ]
    
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="File type not supported")
    
    # Save file
    file_path = os.path.join(UPLOAD_DIR, f"{notebook_id}_{file.filename}")
    content = await file.read()
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Create document record
    db_document = Document(
        notebook_id=notebook_id,
        filename=file.filename,
        file_type=file.content_type,
        file_url=file_path,
        file_size=len(content),
        status="pending"
    )
    
    db.add(db_document)
    db.commit()
    db.refresh(db_document)
    
    logger.info(f"Document {db_document.id} uploaded successfully: {file.filename}")
    
    # Start asynchronous processing with a new database session
    asyncio.create_task(process_document(db_document.id, file_path, file.content_type))
    
    return db_document

@router.delete("/{document_id}")
def delete_document(document_id: int, db: Session = Depends(get_db)):
    """Delete a document"""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete file
    if os.path.exists(document.file_url):
        os.remove(document.file_url)
    
    db.delete(document)
    db.commit()
    
    return {"message": "Document deleted successfully"}

@router.get("/preview/{document_id}")
def preview_document(document_id: int, db: Session = Depends(get_db)):
    """Preview a document - returns the file for inline viewing"""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not os.path.exists(document.file_url):
        raise HTTPException(status_code=404, detail="File not found")
    
    filename = document.filename
    file_type = document.file_type
    encoded_filename = quote(filename, safe='')
    
    if file_type == "application/pdf":
        return FileResponse(
            path=document.file_url,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"inline; filename*=UTF-8''{encoded_filename}"
            }
        )
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return FileResponse(
            path=document.file_url,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"inline; filename*=UTF-8''{encoded_filename}"
            }
        )
    elif file_type in ["text/plain", "text/markdown"]:
        with open(document.file_url, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return Response(
            content=content,
            media_type="text/plain; charset=utf-8",
            headers={
                "Content-Disposition": f"inline; filename*=UTF-8''{encoded_filename}"
            }
        )
    elif file_type == "text/html":
        with open(document.file_url, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return Response(
            content=content,
            media_type="text/html; charset=utf-8",
            headers={
                "Content-Disposition": f"inline; filename*=UTF-8''{encoded_filename}"
            }
        )
    else:
        return FileResponse(
            path=document.file_url,
            media_type=file_type,
            headers={
                "Content-Disposition": f"inline; filename*=UTF-8''{encoded_filename}"
            }
        )
